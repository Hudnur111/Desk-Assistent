import asyncio
from pathlib import Path

from docx import Document

from .base import Tool

DOCUMENTS_DIR = Path("data/documents")


def _safe_path(title: str) -> Path:
    DOCUMENTS_DIR.mkdir(parents=True, exist_ok=True)
    return DOCUMENTS_DIR / f"{title}.docx"


async def _create_document(title: str, content: str) -> str:
    def _write() -> Path:
        path = _safe_path(title)
        document = Document()
        document.add_heading(title, level=1)
        for paragraph in content.split("\n\n"):
            document.add_paragraph(paragraph)
        document.save(path)
        return path

    path = await asyncio.to_thread(_write)
    return f"Dokument erstellt: {path}"


async def _append_to_document(path: str, text: str) -> str:
    file_path = Path(path)
    if not file_path.is_file():
        return f"Dokument nicht gefunden: {path}"

    def _write() -> None:
        document = Document(file_path)
        document.add_paragraph(text)
        document.save(file_path)

    await asyncio.to_thread(_write)
    return f"Text an {path} angehaengt"


async def _read_document(path: str) -> str:
    file_path = Path(path)

    def _read() -> str | None:
        if not file_path.is_file():
            return None
        document = Document(file_path)
        return "\n".join(p.text for p in document.paragraphs)

    content = await asyncio.to_thread(_read)
    if content is None:
        return f"Dokument nicht gefunden: {path}"
    return content


def document_tools() -> list[Tool]:
    return [
        Tool(
            name="create_document",
            description="Erstellt ein neues Word-Dokument (.docx) mit Titel und Inhalt.",
            input_schema={
                "type": "object",
                "properties": {
                    "title": {
                        "type": "string",
                        "description": "Titel des Dokuments, wird auch als Dateiname verwendet",
                    },
                    "content": {
                        "type": "string",
                        "description": "Inhalt; durch Leerzeilen getrennte Abschnitte werden als eigene Absaetze angelegt",
                    },
                },
                "required": ["title", "content"],
            },
            handler=_create_document,
        ),
        Tool(
            name="append_to_document",
            description="Haengt einen weiteren Absatz an ein bestehendes Word-Dokument an.",
            input_schema={
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Pfad zur .docx-Datei"},
                    "text": {"type": "string", "description": "Anzuhaengender Absatz"},
                },
                "required": ["path", "text"],
            },
            handler=_append_to_document,
        ),
        Tool(
            name="read_document",
            description="Liest den Textinhalt eines Word-Dokuments (.docx).",
            input_schema={
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Pfad zur .docx-Datei"}
                },
                "required": ["path"],
            },
            handler=_read_document,
        ),
    ]
